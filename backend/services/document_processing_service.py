import hashlib
import ipaddress
import logging
import mimetypes
import os
import re
import socket
import time
from contextlib import contextmanager
from dataclasses import asdict
from datetime import datetime
from pathlib import Path
from urllib.parse import urlparse
from uuid import uuid4

import httpx
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from fastapi import HTTPException, UploadFile
from pypdf import PdfReader
from sqlalchemy.orm import Session

from database.connection import BACKEND_DIR, SessionLocal
from database.models import Bot, Chunk, Document, IngestionJob, Website, WebsiteCrawl
from services.chunking_service import (
    CHUNK_OVERLAP,
    CHUNK_SIZE,
    chunk_text_with_metadata,
    count_tokens,
    normalize_text,
)
from services.embedding_service import (
    EMBEDDING_BATCH_SIZE,
    generate_embeddings_batch,
    get_last_embedding_metadata,
)
from services.crawler_service import CrawlAuditReport, get_crawler_provider
from services.firecrawl_service import normalize_crawl_url
from services.object_storage import (
    ObjectStorageError,
    build_source_object_key,
    get_object_storage,
    validate_source_object_ownership,
)
from services.coverage_manifest_service import (
    build_website_coverage_manifest,
    infer_document_relationships,
)
from services.usage_service import ensure_can_promote_knowledge


UPLOAD_DIR = Path(os.getenv("KNOWLEDGE_UPLOAD_DIR", BACKEND_DIR / "storage" / "knowledge")).resolve()
MAX_UPLOAD_BYTES = int(os.getenv("KNOWLEDGE_MAX_UPLOAD_MB", "20")) * 1024 * 1024
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".docx"}
SUPPORTED_MIME_TYPES = {
    "application/pdf": ".pdf",
    "text/plain": ".txt",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": ".docx",
    "application/octet-stream": None,
}
REQUEST_HEADERS = {
    "User-Agent": (
        "ChatbotSaaSKnowledgeBot/4.0 "
        "(public website ingestion; +https://example.com/bot) "
        "Mozilla/5.0"
    ),
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
    "Referer": "https://www.google.com/",
}
REQUEST_TIMEOUT = httpx.Timeout(30.0, connect=10.0)
RETRY_STATUS_CODES = {408, 429, 500, 502, 503, 504}
logger = logging.getLogger("backend.knowledge_ingestion")


def serialize_document(document: Document) -> dict:
    website = document.website
    crawl = document.crawl
    active = document.status == "ready" and (
        website is None or (document.crawl_id is not None and document.crawl_id == website.active_crawl_id)
    )
    lifecycle = {
        "staging": "processing",
        "processing_failed": "failed",
        "stale": "superseded",
        "ready": "active" if active else "ready",
    }.get(document.status, document.status or "ready")
    safe_processing_error = None
    if document.processing_error:
        from workers.job_models import sanitize_customer_error
        _, safe_processing_error = sanitize_customer_error(Exception(document.processing_error))
    page_count = 1
    logical_size_bytes = document.logical_size_bytes
    chunk_count = document.chunk_count
    token_count = document.token_count
    if crawl:
        active_pages = [
            item for item in (website.documents if website else [])
            if item.crawl_id == crawl.id and item.status == "ready"
        ]
        page_count = len(active_pages) or int((crawl.audit_metadata or {}).get("stored_documents") or crawl.pages_crawled or 0)
        if active_pages:
            logical_size_bytes = sum(int(item.logical_size_bytes or 0) for item in active_pages)
            chunk_count = sum(int(item.chunk_count or 0) for item in active_pages)
            token_count = sum(int(item.token_count or 0) for item in active_pages)
    return {
        "id": document.id,
        "bot_id": document.bot_id,
        "filename": document.filename,
        "source_type": document.source_type,
        "source_url": document.source_url,
        "file_size": document.file_size,
        "logical_size_bytes": logical_size_bytes,
        "processing_status": document.processing_status,
        "processing_error": safe_processing_error,
        "chunk_count": chunk_count,
        "token_count": token_count,
        "lifecycle_status": lifecycle,
        "active": active,
        "version": document.version,
        "page_count": page_count,
        "last_indexed_at": document.last_crawled_at or (crawl.completed_at if crawl else document.updated_at),
        "metadata": document.metadata_json or {},
        "created_at": document.created_at,
        "updated_at": document.updated_at,
    }


def sanitize_filename(filename: str) -> str:
    name = Path(filename or "upload").name
    name = re.sub(r"[^A-Za-z0-9._-]+", "-", name).strip(".-")
    return name[:180] or f"upload-{uuid4().hex}"


def validate_upload(file: UploadFile) -> str:
    filename = sanitize_filename(file.filename or "")
    extension = Path(filename).suffix.lower()
    content_type = (file.content_type or mimetypes.guess_type(filename)[0] or "").split(";", 1)[0].lower().strip()
    mime_extension = SUPPORTED_MIME_TYPES.get(content_type)

    if extension not in SUPPORTED_EXTENSIONS:
        raise HTTPException(status_code=400, detail="Supported file types are PDF, TXT, and DOCX.")
    if content_type and content_type not in SUPPORTED_MIME_TYPES:
        raise HTTPException(status_code=400, detail=f"Unsupported MIME type: {content_type}")
    if mime_extension and mime_extension != extension:
        raise HTTPException(status_code=400, detail="File extension does not match its MIME type.")
    return filename


def _read_validated_upload(file: UploadFile, filename: str) -> bytes:
    """Read a bounded upload and apply lightweight content-signature checks."""
    chunks: list[bytes] = []
    size = 0
    while True:
        chunk = file.file.read(1024 * 1024)
        if not chunk:
            break
        size += len(chunk)
        if size > MAX_UPLOAD_BYTES:
            raise HTTPException(status_code=413, detail="File exceeds upload size limit.")
        chunks.append(chunk)
    data = b"".join(chunks)
    extension = Path(filename).suffix.lower()
    if extension == ".pdf" and not data.startswith(b"%PDF-"):
        raise HTTPException(status_code=400, detail="The uploaded file is not a valid PDF.")
    if extension == ".docx" and not data.startswith(b"PK"):
        raise HTTPException(status_code=400, detail="The uploaded file is not a valid DOCX archive.")
    if extension == ".txt" and b"\x00" in data[:8192]:
        raise HTTPException(status_code=400, detail="The uploaded TXT file appears to contain binary data.")
    return data


def save_upload(
    file: UploadFile,
    bot_id: int,
    organization_id: int | None = None,
) -> tuple[str, int, str, str]:
    """Persist an upload through the configured object-storage adapter."""
    filename = validate_upload(file)
    if organization_id is None:
        raise HTTPException(status_code=409, detail="Organization ownership is required for durable uploads.")
    data = _read_validated_upload(file, filename)
    storage = get_object_storage()
    key = build_source_object_key(organization_id, bot_id, Path(filename).suffix)
    stored = storage.put(
        key,
        data,
        content_type=file.content_type,
        metadata={"organization_id": str(organization_id), "bot_id": str(bot_id)},
    )
    return stored.key, stored.size, storage.provider_name, hashlib.sha256(data).hexdigest()


def validate_public_url(url: str) -> str:
    parsed = urlparse(url)
    if parsed.scheme not in {"http", "https"} or not parsed.netloc:
        raise HTTPException(status_code=400, detail="URL must be an absolute http(s) URL.")

    hostname = parsed.hostname
    if not hostname:
        raise HTTPException(status_code=400, detail="URL hostname is required.")

    try:
        addresses = socket.getaddrinfo(hostname, None)
    except socket.gaierror as exc:
        raise HTTPException(status_code=400, detail="URL hostname could not be resolved.") from exc

    for address in addresses:
        ip = ipaddress.ip_address(address[4][0])
        if ip.is_private or ip.is_loopback or ip.is_link_local or ip.is_multicast or ip.is_reserved:
            raise HTTPException(status_code=400, detail="Private or local URLs are not allowed.")
    return url


def extract_text_from_file(path: str, source_type: str) -> str:
    if source_type == "txt":
        return Path(path).read_text(encoding="utf-8", errors="ignore")
    if source_type == "pdf":
        reader = PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if source_type == "docx":
        doc = DocxDocument(path)
        return "\n".join(paragraph.text for paragraph in doc.paragraphs)
    raise ValueError(f"Unsupported source type: {source_type}")


@contextmanager
def materialize_document_source(document: Document):
    """Yield a secure temporary local copy for an object-backed source."""
    if document.storage_provider and document.storage_key:
        safe_key = validate_source_object_ownership(
            document.storage_key,
            document.organization_id,
            document.bot_id,
        )
        storage = get_object_storage(document.storage_provider)
        with storage.download_to_temp(safe_key) as temporary_path:
            if document.source_content_hash:
                actual_hash = hashlib.sha256(Path(temporary_path).read_bytes()).hexdigest()
                if actual_hash != document.source_content_hash:
                    raise ObjectStorageError("The durable source object failed its integrity check.")
            yield temporary_path
        return
    if document.file_path:
        # Legacy rows remain readable during the additive rollout. New
        # production uploads never create API-local file paths.
        yield document.file_path
        return
    raise ObjectStorageError("The durable uploaded source object is unavailable.")


def fetch_html(url: str) -> httpx.Response:
    normalized_url = validate_public_url(url)
    last_error: Exception | None = None
    with httpx.Client(timeout=REQUEST_TIMEOUT, follow_redirects=True, headers=REQUEST_HEADERS) as client:
        for attempt in range(3):
            try:
                response = client.get(normalized_url)
                if response.status_code not in RETRY_STATUS_CODES:
                    response.raise_for_status()
                    return response
                response.raise_for_status()
            except (httpx.TimeoutException, httpx.TransportError, httpx.HTTPStatusError) as exc:
                last_error = exc
                status_code = exc.response.status_code if isinstance(exc, httpx.HTTPStatusError) else None
                if attempt == 2 or (status_code is not None and status_code not in RETRY_STATUS_CODES):
                    raise
                time.sleep(0.5 * (attempt + 1))
    raise RuntimeError("Unable to fetch URL.") from last_error


def extract_readable_text_from_html(html: str, fallback_title: str) -> tuple[str, str]:
    from services.scraper_service import _extract_text_from_html
    title, content = _extract_text_from_html(html)
    return title or fallback_title, content


def extract_text_from_url(url: str) -> tuple[str, str]:
    response = fetch_html(url)

    content_type = response.headers.get("content-type", "")
    if "text/html" not in content_type and "application/xhtml+xml" not in content_type:
        raise ValueError("URL did not return an HTML page.")

    return extract_readable_text_from_html(response.text, url)


def _get_scoped_bot(db: Session, bot_id: int) -> Bot:
    bot = db.query(Bot).filter(Bot.id == bot_id).first()
    if not bot:
        raise HTTPException(status_code=404, detail="Bot not found")
    if bot.organization_id is None:
        raise HTTPException(
            status_code=409,
            detail="This legacy bot must be assigned to an organization before knowledge can be processed.",
        )
    return bot


def create_file_document(db: Session, bot_id: int, file: UploadFile) -> Document:
    bot = _get_scoped_bot(db, bot_id)
    filename = validate_upload(file)
    storage_key, size, storage_provider, source_content_hash = save_upload(
        file, bot_id, bot.organization_id
    )
    try:
        existing = db.query(Document).filter(
            Document.bot_id == bot_id,
            Document.organization_id == bot.organization_id,
            Document.source_content_hash == source_content_hash,
            Document.status.in_(["staging", "ready"]),
        ).first()
    except Exception:
        try:
            get_object_storage(storage_provider).delete(storage_key)
        except ObjectStorageError:
            logger.warning("Failed to remove an object after duplicate detection failed.")
        raise
    if existing:
        try:
            get_object_storage(storage_provider).delete(storage_key)
        except ObjectStorageError:
            logger.warning("Failed to remove a duplicate upload object.")
        raise HTTPException(status_code=409, detail="This file content already exists in the bot knowledge base.")
    source_type = Path(filename).suffix.lower().lstrip(".")
    document = Document(
        bot_id=bot_id,
        organization_id=bot.organization_id,
        filename=filename,
        source_type=source_type,
        title=filename,
        raw_text="",
        file_path=None,
        file_size=size,
        storage_provider=storage_provider,
        storage_key=storage_key,
        content_type=file.content_type,
        original_filename=file.filename,
        source_content_hash=source_content_hash,
        logical_size_bytes=size,
        processing_status="pending",
        status="staging",
        metadata_json={
            "original_filename": file.filename,
            "content_type": file.content_type,
            "storage_provider": storage_provider,
        },
    )
    try:
        db.add(document)
        db.commit()
        db.refresh(document)
    except Exception:
        db.rollback()
        try:
            get_object_storage(storage_provider).delete(storage_key)
        except ObjectStorageError:
            logger.warning("Failed to remove an object after document persistence failed.")
        raise
    return document


def create_website_document(
    db: Session,
    bot_id: int,
    url: str,
    crawl_mode: str = "recursive",
) -> Document:
    bot = _get_scoped_bot(db, bot_id)
    normalized_url = normalize_crawl_url(validate_public_url(url))
    existing = (
        db.query(Document)
        .filter(Document.bot_id == bot_id, Document.source_url == normalized_url)
        .first()
    )
    if existing:
        raise HTTPException(status_code=409, detail="This URL is already in the bot knowledge base.")

    document = Document(
        bot_id=bot_id,
        organization_id=bot.organization_id,
        filename=urlparse(normalized_url).netloc,
        source_type="website",
        source_url=normalized_url,
        title=normalized_url,
        raw_text="",
        logical_size_bytes=0,
        processing_status="pending",
        status="staging",
        metadata_json={"crawl_mode": crawl_mode},
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    return document


def remove_unreferenced_upload(
    db: Session,
    file_path: str | None,
    storage_provider: str | None = None,
    storage_key: str | None = None,
    organization_id: int | None = None,
    bot_id: int | None = None,
) -> bool:
    """Remove a durable source object only after its final DB reference is gone."""
    if storage_key and storage_provider:
        if organization_id is None or bot_id is None:
            logger.warning("Refusing to delete an object-storage source without tenant scope.")
            return False
        try:
            storage_key = validate_source_object_ownership(storage_key, organization_id, bot_id)
        except ObjectStorageError:
            logger.warning("Refusing to delete an object-storage source outside its tenant scope.")
            return False
        if db.query(Document.id).filter(
            Document.storage_provider == storage_provider,
            Document.storage_key == storage_key,
        ).first():
            return False
        try:
            return get_object_storage(storage_provider).delete(storage_key)
        except ObjectStorageError:
            logger.warning("Failed to remove an unreferenced knowledge source object.")
            return False

    # Backward-compatible cleanup for pre-object-storage local documents.
    if not file_path:
        return False
    candidate = Path(file_path).resolve()
    try:
        candidate.relative_to(UPLOAD_DIR)
    except ValueError:
        return False
    if db.query(Document.id).filter(Document.file_path == str(candidate)).first():
        return False
    if not candidate.is_file():
        return False
    try:
        candidate.unlink()
    except OSError:
        logger.warning("Failed to remove an unreferenced local knowledge upload.")
        return False
    try:
        candidate.parent.rmdir()
    except OSError:
        pass
    return True


class IngestionCancelled(RuntimeError):
    pass


def _crawl_website_for_mode(
    root_url: str,
    crawl_mode: str,
    cancel_check,
):
    if crawl_mode == "single_page":
        return scrape_single_page_with_audit(root_url, cancel_check=cancel_check)
    return crawl_website(root_url, return_audit=True, cancel_check=cancel_check)


# Stable compatibility seams retained for existing deterministic tests and
# internal callers while the processing core itself uses the crawler port.
def scrape_single_page_with_audit(root_url: str, cancel_check=None):
    return get_crawler_provider().fetch_exact_page(root_url, cancel_check=cancel_check)


def crawl_website(root_url: str, return_audit: bool = False, cancel_check=None):
    pages, audit = get_crawler_provider().crawl_site(root_url, cancel_check=cancel_check)
    return (pages, audit) if return_audit else pages


def _job_is_cancelled(db: Session, job_id: str | None) -> bool:
    if not job_id:
        return False
    job = (
        db.query(IngestionJob)
        .filter(IngestionJob.job_id == job_id)
        .populate_existing()
        .first()
    )
    return bool(job and job.status == "cancelled")


def _assert_not_cancelled(db: Session, job_id: str | None) -> None:
    if _job_is_cancelled(db, job_id):
        raise IngestionCancelled("Ingestion job was cancelled.")


def _set_job_stage(db: Session, job_id: str | None, status: str, progress: int) -> None:
    if not job_id:
        return
    from workers.job_models import transition_job_state

    job = db.query(IngestionJob).filter(IngestionJob.job_id == job_id).first()
    if job and job.status != status:
        transition_job_state(db, job_id, status, stage=status, progress_percent=progress)


def _embed_in_cancellable_batches(
    db: Session,
    texts: list[str],
    organization_id: int,
    job_id: str | None,
) -> tuple[list[list[float]], dict[str, object]]:
    vectors: list[list[float]] = []
    effective_metadata: dict[str, object] | None = None
    for offset in range(0, len(texts), EMBEDDING_BATCH_SIZE):
        _assert_not_cancelled(db, job_id)
        batch = texts[offset : offset + EMBEDDING_BATCH_SIZE]
        vectors.extend(generate_embeddings_batch(batch, org_id=organization_id))
        batch_metadata = get_last_embedding_metadata()
        if effective_metadata and batch_metadata != effective_metadata:
            raise RuntimeError("Embedding provider changed within one ingestion job.")
        effective_metadata = batch_metadata
        if job_id:
            db.query(IngestionJob).filter(IngestionJob.job_id == job_id).update(
                {"last_heartbeat": datetime.utcnow(), "updated_at": datetime.utcnow()}
            )
            db.commit()
    return vectors, effective_metadata or {}


def _mark_job_ready(
    job: IngestionJob | None,
    *,
    documents: int,
    chunks: int,
    embeddings: int,
    audit: dict | None = None,
) -> None:
    if not job:
        return
    job.status = "ready"
    job.current_stage = "ready"
    job.progress_percent = 100
    job.documents_created = documents
    job.chunks_created = chunks
    job.embeddings_created = embeddings
    job.completed_at = datetime.utcnow()
    job.last_heartbeat = datetime.utcnow()
    job.error_code = None
    job.error_message = None
    if audit is not None:
        job.audit_metadata = audit


def _cleanup_failed_staging(
    db: Session,
    staging_id: str,
    root_document_id: int,
    crawl: WebsiteCrawl | None,
    website: Website | None,
    error: Exception,
) -> Document:
    cancelled = isinstance(error, IngestionCancelled)
    db.rollback()
    db.query(Chunk).filter(Chunk.ingestion_job_id == staging_id).delete(synchronize_session=False)
    staged_children = (
        db.query(Document)
        .filter(Document.ingestion_job_id == staging_id, Document.id != root_document_id)
        .all()
    )
    for child in staged_children:
        db.delete(child)

    document = db.query(Document).filter(Document.id == root_document_id).first()
    if document:
        active_chunks = (
            db.query(Chunk)
            .filter(Chunk.document_id == document.id, Chunk.status == "ready")
            .count()
        )
        document.ingestion_job_id = None
        document.processing_error = str(error)[:2000]
        if active_chunks:
            document.status = "ready"
            document.processing_status = "completed"
        else:
            document.status = "processing_failed"
            document.processing_status = "failed"
            document.chunk_count = 0
            document.token_count = 0
        document.updated_at = datetime.utcnow()

    if crawl:
        crawl = db.query(WebsiteCrawl).filter(WebsiteCrawl.id == crawl.id).first()
        if crawl:
            crawl.status = "cancelled" if cancelled else "failed"
            crawl.error_summary = "Crawl cancelled before activation." if cancelled else "Crawl processing failed."
            crawl.completed_at = datetime.utcnow()
    if website:
        website = db.query(Website).filter(Website.id == website.id).first()
        if website:
            website.crawl_status = "cancelled" if cancelled else "failed"
            website.status = "ready" if website.active_crawl_id else "failed"
    job = db.query(IngestionJob).filter(IngestionJob.job_id == staging_id).first()
    if job and job.status != "cancelled":
        from workers.job_models import sanitize_customer_error

        error_code, error_message = sanitize_customer_error(error)
        job.status = "failed"
        job.current_stage = "failed"
        job.error_code = error_code
        job.error_message = error_message
        job.completed_at = datetime.utcnow()
        job.last_heartbeat = datetime.utcnow()
    db.commit()
    if document:
        db.refresh(document)
        return document
    raise ValueError(f"Document {root_document_id} was not found during staging cleanup")


def _process_file_document(
    db: Session,
    document: Document,
    staging_id: str,
    job_id: str | None,
) -> Document:
    _assert_not_cancelled(db, job_id)
    if document.source_type == "text" and document.raw_text:
        extracted_text = document.raw_text
    else:
        with materialize_document_source(document) as source_path:
            extracted_text = extract_text_from_file(source_path, document.source_type)
    normalized = normalize_text(extracted_text)
    if not normalized:
        raise ValueError("Document extraction produced no usable text.")

    chunk_specs = chunk_text_with_metadata(
        normalized,
        chunk_size=CHUNK_SIZE,
        overlap=CHUNK_OVERLAP,
        page_title=document.title or document.filename,
        source_url=document.source_url,
        metadata=document.metadata_json or {},
    )
    if not chunk_specs:
        raise ValueError("Document chunking produced no usable chunks.")

    _set_job_stage(db, job_id, "embedding", 60)
    vectors, embedding_metadata = _embed_in_cancellable_batches(
        db,
        [chunk.content for chunk in chunk_specs],
        document.organization_id,
        job_id,
    )
    for spec, vector in zip(chunk_specs, vectors):
        db.add(
            Chunk(
                document_id=document.id,
                bot_id=document.bot_id,
                organization_id=document.organization_id,
                ingestion_job_id=staging_id,
                chunk_index=spec.index,
                content=spec.content,
                content_hash=hashlib.sha256(spec.content.encode("utf-8")).hexdigest(),
                embedding=vector,
                token_count=spec.token_count,
                status="staging",
                embedding_provider=str(embedding_metadata.get("provider") or "unknown"),
                embedding_model=str(embedding_metadata.get("model") or "unknown"),
                embedding_version=int(embedding_metadata.get("version") or 1),
                metadata_json={
                    "start_token": spec.start_token,
                    "end_token": spec.end_token,
                    "source_type": document.source_type,
                    "source_url": document.source_url,
                    "page_title": document.title or document.filename,
                    "heading": spec.heading,
                    "section": spec.section,
                    "document_id": document.id,
                    "chunk_id": spec.index,
                    "embedding_dimensions": embedding_metadata.get("dimensions"),
                    **(document.metadata_json or {}),
                },
            )
        )
    db.commit()

    _assert_not_cancelled(db, job_id)
    _set_job_stage(db, job_id, "validating", 90)
    db.rollback()
    logical_size = document.file_size or len(normalized.encode("utf-8"))
    ensure_can_promote_knowledge(
        db,
        document.organization_id,
        resulting_documents=1,
        resulting_storage_bytes=logical_size,
        replaced_document_ids=[document.id],
    )
    locked_doc = db.query(Document).filter(Document.id == document.id).with_for_update().first()
    locked_job = None
    if job_id:
        locked_job = db.query(IngestionJob).filter(IngestionJob.job_id == job_id).with_for_update().first()
        if not locked_job or locked_job.status == "cancelled":
            raise IngestionCancelled("Ingestion job was cancelled before promotion.")

    had_active = db.query(Chunk).filter(Chunk.document_id == locked_doc.id, Chunk.status == "ready").count() > 0
    db.query(Chunk).filter(
        Chunk.document_id == locked_doc.id,
        Chunk.status == "ready",
        Chunk.ingestion_job_id.is_(None),
    ).update({"status": "stale"}, synchronize_session=False)
    db.query(Chunk).filter(Chunk.ingestion_job_id == staging_id).update(
        {"status": "ready", "ingestion_job_id": None}, synchronize_session=False
    )
    locked_doc.raw_text = normalized
    locked_doc.logical_size_bytes = logical_size
    locked_doc.content_hash = hashlib.sha256(normalized.encode("utf-8")).hexdigest()
    locked_doc.processing_status = "completed"
    locked_doc.status = "ready"
    locked_doc.processing_error = None
    locked_doc.ingestion_job_id = None
    locked_doc.chunk_count = len(chunk_specs)
    locked_doc.token_count = count_tokens(normalized)
    locked_doc.version = locked_doc.version + 1 if had_active else max(locked_doc.version, 1)
    locked_doc.embedding_provider = str(embedding_metadata.get("provider") or "unknown")
    locked_doc.embedding_model = str(embedding_metadata.get("model") or "unknown")
    locked_doc.embedding_version = int(embedding_metadata.get("version") or 1)
    locked_doc.embedding_dimensions = int(embedding_metadata.get("dimensions") or 0) or None
    locked_doc.updated_at = datetime.utcnow()
    _mark_job_ready(
        locked_job,
        documents=1,
        chunks=len(chunk_specs),
        embeddings=len(vectors),
        audit={"embedding": embedding_metadata},
    )
    db.commit()
    db.refresh(locked_doc)
    return locked_doc


def _process_website_document(
    db: Session,
    document: Document,
    staging_id: str,
    job_id: str | None,
) -> Document:
    submitted_root_url = document.source_url or ""
    crawl_mode = str((document.metadata_json or {}).get("crawl_mode") or "recursive")
    if crawl_mode not in {"recursive", "single_page"}:
        crawl_mode = "recursive"
    root_url = normalize_crawl_url(submitted_root_url) or submitted_root_url
    website = (
        db.query(Website)
        .filter(
            Website.bot_id == document.bot_id,
            Website.organization_id == document.organization_id,
            Website.root_url.in_({submitted_root_url, root_url}),
        )
        .with_for_update()
        .first()
    )
    if not website:
        website = Website(
            bot_id=document.bot_id,
            organization_id=document.organization_id,
            root_url=root_url,
            domain=urlparse(root_url).netloc or "unknown-domain",
            status="crawling",
            crawl_status="crawling",
        )
        db.add(website)
        db.flush()
    else:
        website.root_url = root_url
        website.crawl_status = "crawling"
    document.source_url = root_url

    latest_version = (
        db.query(WebsiteCrawl.version)
        .filter(WebsiteCrawl.website_id == website.id)
        .order_by(WebsiteCrawl.version.desc())
        .first()
    )
    crawl = WebsiteCrawl(
        website_id=website.id,
        bot_id=document.bot_id,
        organization_id=document.organization_id,
        version=(int(latest_version[0]) + 1) if latest_version else 1,
        status="processing",
        crawler_provider=get_crawler_provider().provider_name,
        started_at=datetime.utcnow(),
    )
    db.add(crawl)
    db.flush()
    if job_id:
        job = db.query(IngestionJob).filter(IngestionJob.job_id == job_id).first()
        if job:
            job.website_id = website.id
            job.crawl_id = crawl.id
    db.commit()

    try:
        _assert_not_cancelled(db, job_id)
        cancel_check = lambda: _job_is_cancelled(db, job_id)
        crawl_result = _crawl_website_for_mode(root_url, crawl_mode, cancel_check)
        if isinstance(crawl_result, tuple):
            pages, audit_report = crawl_result
        else:
            pages = crawl_result
            audit_report = CrawlAuditReport(
                seed_url=root_url,
                discovered_urls=len(pages),
                eligible_urls=len(pages),
                crawled_urls=len(pages),
                stored_documents=len(pages),
            )
        _assert_not_cancelled(db, job_id)
        if not pages:
            raise ValueError("Website crawl produced no usable pages.")

        page_dicts = [
            {"url": page.url, "title": page.title, "raw_text": page.markdown, "metadata": page.metadata}
            for page in pages
        ]
        nodes = infer_document_relationships(page_dicts, root_url=root_url)
        node_map = {node.url: node for node in nodes}
        page_records: list[dict] = []
        total_chunks = 0
        total_embeddings = 0
        embedding_metadata: dict[str, object] = {}

        _set_job_stage(db, job_id, "embedding", 65)
        for page_index, page in enumerate(pages):
            _assert_not_cancelled(db, job_id)
            markdown = normalize_text(page.markdown)
            if not markdown:
                raise ValueError(f"Crawled page produced no usable text: {page.url}")
            metadata = dict(page.metadata or {})
            metadata["crawl_mode"] = crawl_mode
            node = node_map.get(page.url)
            if node:
                metadata.update(
                    {
                        "parent_url": node.parent_url,
                        "category_path": node.category_path,
                        "entity_type": node.entity_type,
                        "sibling_urls": node.sibling_urls,
                        "hierarchy_depth": node.depth,
                    }
                )
            page_hash = metadata.get("content_hash") or hashlib.sha256(markdown.encode("utf-8")).hexdigest()
            metadata["content_hash"] = page_hash
            target_doc = document if page_index == 0 else (
                db.query(Document)
                .filter(
                    Document.bot_id == document.bot_id,
                    Document.organization_id == document.organization_id,
                    Document.source_url == page.url,
                )
                .first()
            )
            if not target_doc:
                target_doc = Document(
                    bot_id=document.bot_id,
                    organization_id=document.organization_id,
                    website_id=website.id,
                    filename=(page.title[:180] or urlparse(page.url).netloc or "web-page"),
                    source_type="website",
                    source_url=page.url,
                    title=page.title,
                    raw_text="",
                    processing_status="processing",
                    status="staging",
                    ingestion_job_id=staging_id,
                    metadata_json={},
                )
                db.add(target_doc)
                db.flush()

            specs = chunk_text_with_metadata(
                markdown,
                chunk_size=CHUNK_SIZE,
                overlap=CHUNK_OVERLAP,
                page_title=page.title,
                source_url=page.url,
                metadata=metadata,
            )
            if not specs:
                raise ValueError(f"Crawled page produced no chunks: {page.url}")
            vectors, page_embedding_metadata = _embed_in_cancellable_batches(
                db,
                [spec.content for spec in specs],
                document.organization_id,
                job_id,
            )
            if embedding_metadata and page_embedding_metadata != embedding_metadata:
                raise RuntimeError("Embedding provider changed within one website crawl.")
            embedding_metadata = page_embedding_metadata
            for spec, vector in zip(specs, vectors):
                db.add(
                    Chunk(
                        document_id=target_doc.id,
                        bot_id=document.bot_id,
                        organization_id=document.organization_id,
                        website_id=website.id,
                        crawl_id=crawl.id,
                        ingestion_job_id=staging_id,
                        chunk_index=spec.index,
                        content=spec.content,
                        content_hash=hashlib.sha256(spec.content.encode("utf-8")).hexdigest(),
                        embedding=vector,
                        token_count=spec.token_count,
                        status="staging",
                        embedding_provider=str(embedding_metadata.get("provider") or "unknown"),
                        embedding_model=str(embedding_metadata.get("model") or "unknown"),
                        embedding_version=int(embedding_metadata.get("version") or 1),
                        metadata_json={
                            "start_token": spec.start_token,
                            "end_token": spec.end_token,
                            "source_type": "website",
                            "source_url": page.url,
                            "page_title": page.title,
                            "heading": spec.heading,
                            "section": spec.section,
                            "document_id": target_doc.id,
                            "chunk_id": spec.index,
                            "embedding_dimensions": embedding_metadata.get("dimensions"),
                            **metadata,
                        },
                    )
                )
            page_records.append(
                {
                    "document_id": target_doc.id,
                    "url": page.url,
                    "title": page.title,
                    "markdown": markdown,
                    "content_hash": page_hash,
                    "metadata": metadata,
                    "chunk_count": len(specs),
                    "token_count": count_tokens(markdown),
                }
            )
            total_chunks += len(specs)
            total_embeddings += len(vectors)
            db.commit()

        for page_dict, record in zip(page_dicts, page_records):
            page_dict["chunk_count"] = record["chunk_count"]
        manifest = build_website_coverage_manifest(page_dicts, root_url=root_url)
        audit = asdict(audit_report)
        audit.update(
            {
                "stored_documents": len(page_records),
                "chunked_documents": len(page_records),
                "total_chunks": total_chunks,
                "embedded_chunks": total_embeddings,
                "coverage_manifest": manifest,
                "embedding": embedding_metadata,
            }
        )
        eligible = max(int(audit_report.eligible_urls), len(page_records))
        coverage_percent = round((len(page_records) / eligible) * 100.0, 2) if eligible else 0.0
        crawl = db.query(WebsiteCrawl).filter(WebsiteCrawl.id == crawl.id).first()
        crawl.pages_discovered = audit_report.discovered_urls
        crawl.pages_eligible = audit_report.eligible_urls
        crawl.pages_crawled = audit_report.crawled_urls
        crawl.pages_skipped = len(audit_report.skipped_urls)
        crawl.pages_failed = len(audit_report.failed_urls)
        crawl.duplicate_urls_removed = audit_report.duplicate_urls_removed
        crawl.max_depth_reached = audit_report.max_depth_reached
        crawl.coverage_percent = coverage_percent
        crawl.chunks_created = total_chunks
        crawl.embeddings_created = total_embeddings
        crawl.audit_metadata = audit
        crawl.embedding_provider = str(embedding_metadata.get("provider") or "unknown")
        crawl.embedding_model = str(embedding_metadata.get("model") or "unknown")
        crawl.embedding_version = int(embedding_metadata.get("version") or 1)
        crawl.embedding_dimensions = int(embedding_metadata.get("dimensions") or 0) or None
        if job_id:
            job = db.query(IngestionJob).filter(IngestionJob.job_id == job_id).first()
            if job:
                job.pages_discovered = audit_report.discovered_urls
                job.pages_crawled = audit_report.crawled_urls
                job.pages_failed = len(audit_report.failed_urls)
                job.audit_metadata = audit
        db.commit()

        _assert_not_cancelled(db, job_id)
        _set_job_stage(db, job_id, "validating", 90)
        db.rollback()
        ensure_can_promote_knowledge(
            db,
            document.organization_id,
            resulting_documents=len(page_records),
            resulting_storage_bytes=sum(len(record["markdown"].encode("utf-8")) for record in page_records),
            replaced_website_id=website.id,
        )
        locked_website = db.query(Website).filter(Website.id == website.id).with_for_update().first()
        locked_crawl = db.query(WebsiteCrawl).filter(WebsiteCrawl.id == crawl.id).with_for_update().first()
        locked_job = None
        if job_id:
            locked_job = db.query(IngestionJob).filter(IngestionJob.job_id == job_id).with_for_update().first()
            if not locked_job or locked_job.status == "cancelled":
                raise IngestionCancelled("Ingestion job was cancelled before promotion.")

        new_document_ids = [record["document_id"] for record in page_records]
        stale_documents = (
            db.query(Document)
            .filter(
                Document.website_id == locked_website.id,
                Document.status == "ready",
                Document.id.notin_(new_document_ids),
            )
            .all()
        )
        for stale_document in stale_documents:
            stale_document.status = "stale"
            db.query(Chunk).filter(
                Chunk.document_id == stale_document.id,
                Chunk.status == "ready",
            ).update({"status": "stale"}, synchronize_session=False)

        for record in page_records:
            target_doc = db.query(Document).filter(Document.id == record["document_id"]).with_for_update().first()
            db.query(Chunk).filter(
                Chunk.document_id == target_doc.id,
                Chunk.status == "ready",
            ).update({"status": "stale"}, synchronize_session=False)
            db.query(Chunk).filter(
                Chunk.ingestion_job_id == staging_id,
                Chunk.document_id == target_doc.id,
            ).update({"status": "ready", "ingestion_job_id": None}, synchronize_session=False)
            target_doc.website_id = locked_website.id
            target_doc.crawl_id = locked_crawl.id
            target_doc.ingestion_job_id = None
            target_doc.source_url = record["url"]
            target_doc.canonical_url = record["metadata"].get("canonical_url") or record["metadata"].get("canonicalURL")
            target_doc.filename = record["title"][:180] or urlparse(record["url"]).netloc or "web-page"
            target_doc.title = record["title"]
            target_doc.raw_text = record["markdown"]
            target_doc.logical_size_bytes = len(record["markdown"].encode("utf-8"))
            target_doc.content_hash = record["content_hash"]
            target_doc.crawl_depth = int(record["metadata"].get("hierarchy_depth", record["metadata"].get("depth", 0)))
            target_doc.metadata_json = record["metadata"]
            target_doc.processing_status = "completed"
            target_doc.status = "ready"
            target_doc.processing_error = None
            target_doc.chunk_count = record["chunk_count"]
            target_doc.token_count = record["token_count"]
            target_doc.version = locked_crawl.version
            target_doc.embedding_provider = str(embedding_metadata.get("provider") or "unknown")
            target_doc.embedding_model = str(embedding_metadata.get("model") or "unknown")
            target_doc.embedding_version = int(embedding_metadata.get("version") or 1)
            target_doc.embedding_dimensions = int(embedding_metadata.get("dimensions") or 0) or None
            target_doc.last_seen_at = datetime.utcnow()
            target_doc.last_crawled_at = datetime.utcnow()
            target_doc.updated_at = datetime.utcnow()

        previous_crawl_id = locked_website.active_crawl_id
        if previous_crawl_id and previous_crawl_id != locked_crawl.id:
            previous = db.query(WebsiteCrawl).filter(WebsiteCrawl.id == previous_crawl_id).with_for_update().first()
            if previous:
                previous.status = "superseded"
        locked_crawl.status = "ready"
        locked_crawl.completed_at = datetime.utcnow()
        locked_website.active_crawl_id = locked_crawl.id
        locked_website.status = "ready"
        locked_website.crawl_status = "ready"
        locked_website.last_crawled_at = datetime.utcnow()
        _mark_job_ready(
            locked_job,
            documents=len(page_records),
            chunks=total_chunks,
            embeddings=total_embeddings,
            audit=audit,
        )
        db.commit()
        promoted = db.query(Document).filter(Document.id == document.id).first()
        db.refresh(promoted)
        return promoted
    except Exception as exc:
        return _cleanup_failed_staging(db, staging_id, document.id, crawl, website, exc)


def process_document(db: Session, document_id: int, job_id: str | None = None) -> Document:
    """Builds replacement knowledge in staging and promotes it in one short transaction."""
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise ValueError(f"Document {document_id} was not found")
    bot = _get_scoped_bot(db, document.bot_id)
    if document.organization_id != bot.organization_id:
        raise ValueError(f"Document {document_id} does not match its bot organization")

    staging_id = job_id or f"inline_{uuid4().hex}"
    db.query(Chunk).filter(Chunk.ingestion_job_id == staging_id).delete(synchronize_session=False)
    if document.status not in {"ready", "stale"}:
        document.status = "staging"
        document.processing_status = "processing"
        document.ingestion_job_id = staging_id
    document.processing_error = None
    document.updated_at = datetime.utcnow()
    db.commit()

    try:
        if document.source_type == "website":
            processed = _process_website_document(db, document, staging_id, job_id)
        else:
            processed = _process_file_document(db, document, staging_id, job_id)
    except IngestionCancelled as exc:
        processed = _cleanup_failed_staging(db, staging_id, document.id, None, None, exc)
    except Exception as exc:
        processed = _cleanup_failed_staging(db, staging_id, document.id, None, None, exc)

    if processed.status == "ready":
        try:
            from services.rag_service import clear_retrieval_cache
            clear_retrieval_cache(processed.bot_id)
        except Exception:
            pass
    return processed


def process_document_job(document_id: int) -> None:
    db = SessionLocal()
    try:
        process_document(db, document_id)
    finally:
        db.close()


def reindex_document(db: Session, document_id: int) -> Document:
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found.")
    return process_document(db, document_id)
